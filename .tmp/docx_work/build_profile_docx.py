from __future__ import annotations

import json
import math
import re
from datetime import date, timedelta
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
WORK = Path(__file__).resolve().parent
MD_PATH = ROOT / 'documentacion-grado/06_entregables/perfil/PERFIL_PROYECTO_GRADO.md'
OUT_PATH = ROOT / 'documentacion-grado/06_entregables/perfil/Perfil_Proyecto_de_Grado_R3Foresta.docx'
LOGO_PATH = WORK / 'media/image1.png'
PAGE_MAP_PATH = WORK / 'page_map.json'
FIG_DIR = WORK / 'figures'
FIG_DIR.mkdir(exist_ok=True)

FONT_REGULAR = '/System/Library/Fonts/Supplemental/Arial.ttf'
FONT_BOLD = '/System/Library/Fonts/Supplemental/Arial Bold.ttf'
INK = '#1f2937'
LINE = '#64748b'
FILL = '#f1f5f9'
FILL_DARK = '#e2e8f0'
ACCENT = '#30445f'
WHITE = '#ffffff'


def pil_font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def wrap_text(draw, text: str, font, max_width: int):
    words = text.split()
    lines, current = [], ''
    for word in words:
        probe = f'{current} {word}'.strip()
        if current and draw.textbbox((0, 0), probe, font=font)[2] > max_width:
            lines.append(current)
            current = word
        else:
            current = probe
    if current:
        lines.append(current)
    return lines


def draw_centered_text(draw, box, text, font, fill=INK, spacing=4):
    x1, y1, x2, y2 = box
    lines = []
    for part in text.split('\n'):
        lines.extend(wrap_text(draw, part, font, int((x2 - x1) * 0.88)))
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total = sum(heights) + spacing * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += h + spacing


def rounded_box(draw, box, text, font, fill=FILL, outline=LINE, width=3, radius=22):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    draw_centered_text(draw, box, text, font)


def arrow(draw, start, end, fill=LINE, width=4, dashed=False):
    x1, y1 = start
    x2, y2 = end
    if dashed:
        steps = 18
        for i in range(steps):
            if i % 2 == 0:
                a = i / steps
                b = min(1, (i + 1) / steps)
                draw.line((x1 + (x2-x1)*a, y1 + (y2-y1)*a, x1 + (x2-x1)*b, y1 + (y2-y1)*b), fill=fill, width=width)
    else:
        draw.line((x1, y1, x2, y2), fill=fill, width=width)
    ang = math.atan2(y2-y1, x2-x1)
    size = 15
    pts = [(x2, y2), (x2-size*math.cos(ang-0.55), y2-size*math.sin(ang-0.55)), (x2-size*math.cos(ang+0.55), y2-size*math.sin(ang+0.55))]
    draw.polygon(pts, fill=fill)


def make_figures():
    # Figure 1: portrait cause/effect tree for readable print scaling.
    img = Image.new('RGB', (1600, 1900), WHITE)
    d = ImageDraw.Draw(img)
    f = pil_font(34)
    fb = pil_font(34, True)
    effects = [
        'Mayor tiempo y menor completitud al reconstruir el recorrido del material vegetal',
        'Riesgo de inconsistencias en cantidades, saldos y transferencias',
        'Decisiones operativas basadas en información incompleta',
        'Menor capacidad para respaldar la información comunicada a patrocinadores y aliados',
    ]
    causes = [
        'Registros dispersos en fotos, mensajería, cuadernos, redes sociales y memoria',
        'Ausencia de identificadores y relaciones comunes entre etapas',
        'Cambios de ubicación o responsable sin un registro uniforme de procedencia',
        'Transferencias y saldos controlados mediante operaciones separadas',
        'Evidencia fotográfica, temporal o geográfica desvinculada del hecho',
    ]
    effect_boxes = [(70, 70, 760, 250), (840, 70, 1530, 250), (70, 310, 760, 490), (840, 310, 1530, 490)]
    for box, txt in zip(effect_boxes, effects):
        rounded_box(d, box, txt, f)
    center = (150, 700, 1450, 1130)
    d.ellipse(center, fill=FILL_DARK, outline=ACCENT, width=5)
    draw_centered_text(d, center, 'PROBLEMA CENTRAL\n\nLa información sobre el recorrido del material vegetal no se encuentra integrada bajo una cadena de custodia reconstruible con evidencia contrastable', fb)
    for box in effect_boxes:
        arrow(d, ((box[0]+box[2])//2, center[1]), ((box[0]+box[2])//2, box[3]))
    cause_boxes = [(50, 1370, 500, 1600), (575, 1370, 1025, 1600), (1100, 1370, 1550, 1600), (300, 1660, 760, 1870), (840, 1660, 1300, 1870)]
    for box, txt in zip(cause_boxes, causes):
        rounded_box(d, box, txt, pil_font(29))
        arrow(d, ((box[0]+box[2])//2, box[1]), ((box[0]+box[2])//2, center[3]))
    img.save(FIG_DIR / 'figura_1.png', dpi=(220, 220))

    # Figure 2: principal route.
    img = Image.new('RGB', (2000, 760), WHITE)
    d = ImageDraw.Draw(img)
    f = pil_font(32)
    boxes = {
        'r': (60, 410, 500, 610),
        'x': (60, 90, 500, 290),
        'v': (720, 330, 1140, 550),
        'a': (1270, 330, 1580, 550),
        'p': (1660, 250, 1950, 610),
    }
    labels = {
        'r': 'M1 · Recolección\nlote de origen',
        'x': 'Ingreso externo\ncompra o recepción',
        'v': 'M2 · Vivero\neventos y saldo vivo',
        'a': 'Asignación a\nsubcampaña',
        'p': 'M3 · Plantación\nconsumo, ubicación\ny evidencia',
    }
    for k in boxes:
        rounded_box(d, boxes[k], labels[k], f)
    arrow(d, (boxes['r'][2], 510), (boxes['v'][0], 440))
    arrow(d, (boxes['x'][2], 190), (boxes['v'][0], 370), dashed=True)
    arrow(d, (boxes['x'][2], 190), (boxes['p'][0], 300), dashed=True)
    arrow(d, (boxes['v'][2], 440), (boxes['a'][0], 440))
    arrow(d, (boxes['a'][2], 440), (boxes['p'][0], 440))
    img.save(FIG_DIR / 'figura_2.png', dpi=(220, 220))

    # Figure 3: events, rules, and reconstruction.
    img = Image.new('RGB', (2000, 850), WHITE)
    d = ImageDraw.Draw(img)
    f = pil_font(31)
    boxes = {
        'o': (40, 330, 360, 530),
        'v': (460, 270, 830, 590),
        'e': (980, 40, 1410, 220),
        's': (980, 335, 1410, 515),
        'd': (980, 630, 1410, 810),
        'q': (1570, 270, 1960, 590),
    }
    labels = {
        'o': 'Hecho registrado', 'v': 'Reglas de trazabilidad\ne integridad',
        'e': 'Historial de eventos', 's': 'Cantidades y saldos\ncoherentes',
        'd': 'Responsable, tiempo,\nubicación y evidencia', 'q': 'Reconstrucción\nde la traza',
    }
    for k in boxes:
        rounded_box(d, boxes[k], labels[k], f, fill=FILL_DARK if k in ('v','q') else FILL, width=5 if k in ('v','q') else 3)
    arrow(d, (boxes['o'][2], 430), (boxes['v'][0], 430))
    for k, y in [('e', 130), ('s', 425), ('d', 720)]:
        arrow(d, (boxes['v'][2], 430), (boxes[k][0], y))
        arrow(d, (boxes[k][2], y), (boxes['q'][0], 430))
    img.save(FIG_DIR / 'figura_3.png', dpi=(220, 220))

    # Figure 4: Gantt chart matching the phases and iterations in the source.
    img = Image.new('RGB', (2200, 1320), WHITE)
    d = ImageDraw.Draw(img)
    title_font = pil_font(42, True)
    label_font = pil_font(25)
    small_font = pil_font(21)
    d.text((1100, 28), 'Perfil y Proyecto de Grado — julio a noviembre de 2026', font=title_font, fill=INK, anchor='ma')
    tasks = [
        ('Inicio · IN-1', date(2026,7,6), date(2026,7,20)),
        ('Elaboración · EL-1', date(2026,7,20), date(2026,8,17)),
        ('Construcción · CO-1 / M1', date(2026,8,17), date(2026,9,7)),
        ('Construcción · CO-2 / M2', date(2026,9,7), date(2026,9,28)),
        ('Construcción · CO-3 / M3', date(2026,9,28), date(2026,10,19)),
        ('Construcción · CO-4 / integración', date(2026,10,19), date(2026,11,2)),
        ('Transición · TR-1 / pruebas', date(2026,11,2), date(2026,11,9)),
        ('Transición · TR-1 / cierre', date(2026,11,9), date(2026,11,16)),
        ('Especificaciones, pruebas y documentación', date(2026,7,6), date(2026,11,16)),
    ]
    start, end = date(2026,7,6), date(2026,11,16)
    x0, x1 = 740, 2110
    y0, row_h = 160, 112
    days = (end-start).days
    for week in range(0, days+1, 7):
        x = x0 + (x1-x0)*week/days
        d.line((x,y0-35,x,y0+row_h*len(tasks)), fill='#cbd5e1', width=2)
        dt = start + timedelta(days=week)
        d.text((x, y0-58), dt.strftime('%d/%m'), font=small_font, fill=INK, anchor='ms')
    for i,(label, s, e) in enumerate(tasks):
        y = y0 + i*row_h
        d.text((700,y+40), label, font=label_font, fill=INK, anchor='ra')
        xs = x0 + (x1-x0)*(s-start).days/days
        xe = x0 + (x1-x0)*(e-start).days/days
        color = ACCENT if i < 8 else '#94a3b8'
        d.rounded_rectangle((xs,y+10,xe,y+72), radius=10, fill=color, outline=INK, width=2)
    d.line((x1,y0-45,x1,y0+row_h*len(tasks)), fill=INK, width=4)
    img.save(FIG_DIR / 'figura_4.png', dpi=(220, 220))


def set_font(run, name='Times New Roman', size=12, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, italic=False, bold=False):
    part = paragraph.part
    rel_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rel_id)
    new_run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')
    rpr.append(rfonts)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rpr.append(color)
    underline = OxmlElement('w:u'); underline.set(qn('w:val'), 'single'); rpr.append(underline)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24'); rpr.append(sz)
    if italic:
        rpr.append(OxmlElement('w:i'))
    if bold:
        rpr.append(OxmlElement('w:b'))
    new_run.append(rpr)
    text_el = OxmlElement('w:t'); text_el.text = text; new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r'(https?://[^\s)]+|\*\*.+?\*\*|__.+?__|(?<!\*)\*[^*]+?\*|_[^_]+?_)')


def add_inline(paragraph, text: str, size=12):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            set_font(paragraph.add_run(text[pos:m.start()]), size=size)
        token = m.group(0)
        if token.startswith('http'):
            add_hyperlink(paragraph, token, token)
        elif token.startswith(('**','__')):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True)
        else:
            set_font(paragraph.add_run(token[1:-1]), size=size, italic=True)
        pos = m.end()
    if pos < len(text):
        set_font(paragraph.add_run(text[pos:]), size=size)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run, size=10)
    fld_char1 = OxmlElement('w:fldChar'); fld_char1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    fld_char2 = OxmlElement('w:fldChar'); fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_char1, instr, fld_char2])


def set_page_num_type(section, fmt='decimal', start=None):
    sect_pr = section._sectPr
    old = sect_pr.find(qn('w:pgNumType'))
    if old is not None:
        sect_pr.remove(old)
    el = OxmlElement('w:pgNumType')
    el.set(qn('w:fmt'), fmt)
    if start is not None:
        el.set(qn('w:start'), str(start))
    sect_pr.append(el)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar'); tc_pr.append(tc_mar)
    for m, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}'); tc_mar.append(node)
        node.set(qn('w:w'), str(val)); node.set(qn('w:type'),'dxa')


def set_table_borders(table, color='94A3B8', size='4'):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders'); tbl_pr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = borders.find(qn(f'w:{edge}'))
        if tag is None:
            tag = OxmlElement(f'w:{edge}'); borders.append(tag)
        tag.set(qn('w:val'),'single'); tag.set(qn('w:sz'),size); tag.set(qn('w:color'),color)


def set_table_widths(table, widths_twips):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW'); tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths_twips))); tbl_w.set(qn('w:type'),'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement('w:gridCol'); col.set(qn('w:w'),str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW'); tc_pr.append(tc_w)
            tc_w.set(qn('w:w'),str(width)); tc_w.set(qn('w:type'),'dxa')


def add_table(doc, rows, aligns, table_index):
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if cols == 4:
        widths_by_table = {
            1: [1450, 1100, 3000, 2720],
            3: [1150, 2500, 2500, 2120],
        }
        widths = widths_by_table.get(table_index, [1500, 2200, 2400, 2170])
    elif cols == 3:
        widths_by_table = {
            1: [1800, 3200, 3270],
            2: [1800, 3000, 3470],
            3: [1800, 4500, 1970],
            4: [1700, 3900, 2670],
        }
        widths = widths_by_table.get(table_index, [1900, 3300, 3070])
    else:
        widths = [8270//cols]*cols; widths[-1] += 8270-sum(widths)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i,j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 or aligns[j] == 'left' else (WD_ALIGN_PARAGRAPH.RIGHT if aligns[j]=='right' else WD_ALIGN_PARAGRAPH.CENTER)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.right_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value.strip().replace('**',''), size=9)
            if i == 0:
                set_cell_shading(cell, 'DCE6F1')
                for r in p.runs:
                    r.bold = True
            elif i % 2 == 0:
                set_cell_shading(cell, 'F8FAFC')
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader'); tbl_header.set(qn('w:val'),'true'); tr_pr.append(tbl_header)
    set_table_widths(table,widths)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def new_numbering_id(doc, ordered=True):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement('w:abstractNum'); abstract.set(qn('w:abstractNumId'), str(abstract_id))
    multi = OxmlElement('w:multiLevelType'); multi.set(qn('w:val'), 'singleLevel'); abstract.append(multi)
    lvl = OxmlElement('w:lvl'); lvl.set(qn('w:ilvl'), '0'); abstract.append(lvl)
    start = OxmlElement('w:start'); start.set(qn('w:val'), '1'); lvl.append(start)
    fmt = OxmlElement('w:numFmt'); fmt.set(qn('w:val'), 'decimal' if ordered else 'bullet'); lvl.append(fmt)
    txt = OxmlElement('w:lvlText'); txt.set(qn('w:val'), '%1.' if ordered else '•'); lvl.append(txt)
    jc = OxmlElement('w:lvlJc'); jc.set(qn('w:val'), 'left'); lvl.append(jc)
    ppr = OxmlElement('w:pPr'); lvl.append(ppr)
    tabs = OxmlElement('w:tabs'); ppr.append(tabs)
    tab = OxmlElement('w:tab'); tab.set(qn('w:val'), 'num'); tab.set(qn('w:pos'), '720'); tabs.append(tab)
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '720'); ind.set(qn('w:hanging'), '360'); ppr.append(ind)
    if not ordered:
        rpr = OxmlElement('w:rPr'); lvl.append(rpr)
        fonts = OxmlElement('w:rFonts'); fonts.set(qn('w:ascii'), 'Times New Roman'); fonts.set(qn('w:hAnsi'), 'Times New Roman'); rpr.append(fonts)
    numbering.append(abstract)
    num = OxmlElement('w:num'); num.set(qn('w:numId'), str(num_id))
    abstract_ref = OxmlElement('w:abstractNumId'); abstract_ref.set(qn('w:val'), str(abstract_id)); num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn('w:numPr'))
    if num_pr is None:
        num_pr = OxmlElement('w:numPr'); ppr.append(num_pr)
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0'); num_pr.append(ilvl)
    num = OxmlElement('w:numId'); num.set(qn('w:val'), str(num_id)); num_pr.append(num)


def setup_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'; normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman')
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.27)
    for name,size,before,after in [('Heading 1',14,12,6),('Heading 2',12,10,4),('Heading 3',12,8,3)]:
        st=styles[name]; st.font.name='Times New Roman'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor(0,0,0)
        st._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman')
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
        st.paragraph_format.line_spacing=1.15; st.paragraph_format.first_line_indent=Cm(0)
    # Main sections flow continuously; page breaks are reserved for the cover,
    # front matter, and cases where natural pagination requires them.
    styles['Heading 1'].paragraph_format.page_break_before = False
    for name in ['List Bullet','List Number']:
        st=styles[name]; st.font.name='Times New Roman'; st.font.size=Pt(12)
        st._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman')
        st.paragraph_format.line_spacing=1.5; st.paragraph_format.space_after=Pt(0)
    cap=styles['Caption']; cap.font.name='Times New Roman'; cap.font.size=Pt(11); cap.font.bold=True; cap.font.italic=False; cap.font.color.rgb=RGBColor(0,0,0)
    cap._element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman'); cap._element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman')
    cap.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT; cap.paragraph_format.space_before=Pt(6); cap.paragraph_format.space_after=Pt(2); cap.paragraph_format.keep_with_next=True; cap.paragraph_format.line_spacing=1.0; cap.paragraph_format.first_line_indent=Cm(0)


def add_cover(doc):
    sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11)
    sec.left_margin=Cm(4); sec.right_margin=Cm(3); sec.top_margin=Cm(3); sec.bottom_margin=Cm(3)
    def cp(text='', size=12, bold=False, before=0, after=0):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
        if text:
            set_font(p.add_run(text), size=size, bold=bold)
        return p
    cp('UNIVERSIDAD MAYOR DE SAN ANDRÉS',15,True)
    cp('FACULTAD DE CIENCIAS PURAS Y NATURALES',13,True)
    cp('CARRERA DE INFORMÁTICA',13,True,after=5)
    p=cp(after=4)
    picture=p.add_run().add_picture(str(LOGO_PATH),width=Inches(1.6))
    picture._inline.docPr.set('title','Logotipo de R3Foresta')
    picture._inline.docPr.set('descr','Logotipo institucional de R3Foresta')
    cp('PROYECTO DE GRADO',13,True,after=8)
    cp('SISTEMA DE TRAZABILIDAD DEL MATERIAL VEGETAL PARA REFORESTACIÓN',14,True,after=2)
    cp('CON PROYECCIÓN HACIA BONOS DE CARBONO',13,True,after=3)
    cp('CASO: R3FORESTA',13,True,after=8)
    cp('Proyecto de Grado para obtener el Título de Licenciatura en Informática',11)
    cp('Mención Ingeniería de Sistemas Informáticos',11,after=14)
    cp('POR: PABLO ANDRES FERNANDEZ CARI',12,True,after=4)
    cp('TUTORA: Ph. D. MARISOL TÉLLEZ RAMÍREZ',12,True,after=18)
    cp('LA PAZ – BOLIVIA',12,True)
    cp('Agosto, 2026',12,True)


def add_header_footer(section, header_text, fmt='decimal', start=None):
    section.header.is_linked_to_previous=False
    section.footer.is_linked_to_previous=False
    hp=section.header.paragraphs[0]; clear_paragraph(hp); hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; hp.paragraph_format.space_after=Pt(0)
    if header_text:
        set_font(hp.add_run(header_text),size=9,color='6B7280')
    fp=section.footer.paragraphs[0]; clear_paragraph(fp); add_page_field(fp)
    set_page_num_type(section,fmt,start)


def parse_table(lines, start):
    raw=[]; i=start
    while i<len(lines) and lines[i].lstrip().startswith('|'):
        raw.append(lines[i].strip()); i+=1
    def cells(line): return [c.strip() for c in line.strip().strip('|').split('|')]
    rows=[cells(x) for x in raw]
    aligns=['left']*len(rows[0])
    if len(rows)>1 and all(re.fullmatch(r':?-{3,}:?',c.replace(' ','')) for c in rows[1]):
        for j,c in enumerate(rows[1]):
            s=c.replace(' ',''); aligns[j]='center' if s.startswith(':') and s.endswith(':') else ('right' if s.endswith(':') else 'left')
        rows.pop(1)
    return rows,aligns,i


def collect_blocks(md_text):
    lines=md_text.splitlines(); blocks=[]; i=0
    in_body=False
    while i<len(lines):
        line=lines[i]
        if line.startswith('## 1. Introducción'):
            in_body=True
        if not in_body:
            i+=1; continue
        if line.startswith('```'):
            lang=line[3:].strip(); buf=[]; i+=1
            while i<len(lines) and not lines[i].startswith('```'):
                buf.append(lines[i]); i+=1
            i+=1; blocks.append(('code',lang,'\n'.join(buf))); continue
        if line.lstrip().startswith('|'):
            rows,aligns,i=parse_table(lines,i); blocks.append(('table',rows,aligns)); continue
        if not line.strip() or line.strip()=='---': i+=1; continue
        m=re.match(r'^(#{2,4})\s+(.+)$',line)
        if m:
            blocks.append(('heading',len(m.group(1))-1,m.group(2).strip())); i+=1; continue
        if re.match(r'^\*Versión de contenido',line): i+=1; continue
        if re.match(r'^\*\*(Tabla|Figura)\s+[A-Z]?\d+\.?\*\*$',line):
            blocks.append(('caption',line.strip().strip('*').rstrip('.'))); i+=1; continue
        image_match = re.match(r'^!\[([^]]*)\]\(([^)]+)\)$', line.strip())
        if image_match:
            blocks.append(('image', image_match.group(1), image_match.group(2))); i+=1; continue
        if line.startswith('> '):
            buf=[]
            while i<len(lines) and lines[i].startswith('> '): buf.append(lines[i][2:]); i+=1
            blocks.append(('quote',' '.join(buf))); continue
        lm=re.match(r'^(\s*)([-+*]|\d+\.)\s+(.+)$',line)
        if lm:
            blocks.append(('list','number' if lm.group(2)[0].isdigit() else 'bullet',len(lm.group(1)),lm.group(3))); i+=1; continue
        # A Markdown paragraph; preserve only explicit same-paragraph continuation lines.
        buf=[line.strip()]; i+=1
        while i<len(lines) and lines[i].strip() and not re.match(r'^(#{2,4})\s+|^```|^\*\*(Tabla|Figura)\s+[A-Z]?\d+\.?\*\*$|^>|^(\s*)([-+*]|\d+\.)\s+|^\|',lines[i]):
            buf.append(lines[i].strip()); i+=1
        blocks.append(('paragraph',' '.join(buf)))
    return blocks


def get_summary(md_text):
    m=re.search(r'^## Resumen\s*\n(.*?)^---\s*$',md_text,re.M|re.S)
    if not m: return []
    return [p.strip() for p in re.split(r'\n\s*\n',m.group(1)) if p.strip()]


def toc_entry(doc,text,page,level=0):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.left_indent=Cm(0.55*level); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.15
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    set_font(p.add_run(text),size=11)
    p.add_run('\t')
    set_font(p.add_run(str(page if page is not None else '00')),size=11)


def build():
    make_figures()
    md=MD_PATH.read_text(encoding='utf-8')
    blocks=collect_blocks(md)
    page_map=json.loads(PAGE_MAP_PATH.read_text()) if PAGE_MAP_PATH.exists() else {}
    doc=Document()
    setup_styles(doc)
    add_cover(doc)
    front=doc.add_section(WD_SECTION.NEW_PAGE)
    front.page_width=Inches(8.5); front.page_height=Inches(11); front.left_margin=Cm(4); front.right_margin=Cm(3); front.top_margin=Cm(3); front.bottom_margin=Cm(3)
    add_header_footer(front,'',fmt='lowerRoman',start=1)
    summaries=get_summary(md)
    if summaries:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_after=Pt(10); set_font(p.add_run('Resumen'),size=14,bold=True)
        for text in summaries:
            p=doc.add_paragraph(); add_inline(p,text); p.paragraph_format.space_after=Pt(6)
            if text.startswith('**Palabras clave:**'):
                p.paragraph_format.first_line_indent=Cm(0)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.page_break_before=bool(summaries); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_after=Pt(10); set_font(p.add_run('Índice general'),size=14,bold=True)
    for block in blocks:
        if block[0]=='heading' and block[1]==1:
            toc_entry(doc,block[2],page_map.get(block[2]))
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.page_break_before=True; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_after=Pt(10); set_font(p.add_run('Índice de tablas'),size=14,bold=True)
    for idx,block in enumerate(blocks):
        if block[0]=='caption' and block[1].startswith('Tabla'):
            title=blocks[idx+1][1].strip('*') if idx+1<len(blocks) and blocks[idx+1][0]=='paragraph' and re.fullmatch(r'\*[^*]+\*',blocks[idx+1][1]) else ''
            toc_entry(doc,f"{block[1]}. {title}".rstrip('. '),page_map.get(block[1]))
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(10); set_font(p.add_run('Índice de figuras'),size=14,bold=True)
    for idx,block in enumerate(blocks):
        if block[0]=='caption' and block[1].startswith('Figura'):
            title=blocks[idx+1][1].strip('*') if idx+1<len(blocks) and blocks[idx+1][0]=='paragraph' and re.fullmatch(r'\*[^*]+\*',blocks[idx+1][1]) else ''
            toc_entry(doc,f"{block[1]}. {title}".rstrip('. '),page_map.get(block[1]))

    body=doc.add_section(WD_SECTION.NEW_PAGE)
    body.page_width=Inches(8.5); body.page_height=Inches(11); body.left_margin=Cm(4); body.right_margin=Cm(3); body.top_margin=Cm(3); body.bottom_margin=Cm(3)
    add_header_footer(body,'Perfil de Proyecto de Grado · R3Foresta',fmt='decimal',start=1)
    current_h1=''
    table_index=0
    previous_kind=''
    list_num_id=None
    list_type=None
    for block in blocks:
        kind=block[0]
        if kind=='heading':
            level,title=block[1],block[2]
            current_h1=title if level==1 else current_h1
            p=doc.add_paragraph(style=f'Heading {level}')
            if level==1 and title=='12. Anexos':
                p.paragraph_format.page_break_before=True
            add_inline(p,title,size=14 if level==1 else 12)
        elif kind=='paragraph':
            text=block[1]
            p=doc.add_paragraph(); add_inline(p,text)
            if current_h1=='11. Referencias bibliográficas':
                p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.first_line_indent=Cm(-1.27); p.paragraph_format.left_indent=Cm(1.27); p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(4)
            elif current_h1=='9. Índice propuesto del Proyecto de Grado':
                p.paragraph_format.first_line_indent=Cm(0)
            if previous_kind=='caption' and re.fullmatch(r'\*[^*]+\*',text):
                p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_after=Pt(4); p.paragraph_format.keep_with_next=True
        elif kind=='list':
            if previous_kind!='list' or list_type!=block[1]:
                list_num_id=new_numbering_id(doc,ordered=block[1]=='number')
                list_type=block[1]
            p=doc.add_paragraph(style='List Paragraph'); add_inline(p,block[3]); apply_numbering(p,list_num_id)
            p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_after=Pt(0); p.paragraph_format.left_indent=Cm(0.8+0.4*(block[2]//2)); p.paragraph_format.first_line_indent=Cm(-0.5)
        elif kind=='quote':
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(1); p.paragraph_format.right_indent=Cm(0.6); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(5)
            add_inline(p,block[1]);
            for r in p.runs: r.italic=True
        elif kind=='caption':
            p=doc.add_paragraph(style='Caption'); add_inline(p,block[1],size=11)
        elif kind=='image':
            image_path = ROOT / 'documentacion-grado/06_entregables/perfil' / block[2]
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_after=Pt(6)
            picture=p.add_run().add_picture(str(image_path), width=Inches(5.72))
            picture._inline.docPr.set('title',block[1])
            picture._inline.docPr.set('descr',block[1])
        elif kind=='table':
            table_index+=1
            add_table(doc,block[1],block[2],table_index)
        elif kind=='code':
            lang,content=block[1],block[2]
            if lang=='mermaid':
                p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.space_after=Pt(6)
                width=Inches(5.72)
                is_gantt=content.lstrip().startswith('gantt')
                figure_file='figura_4.png' if is_gantt else 'figura_1.png'
                alt_text=('Diagrama de Gantt del proceso de desarrollo, julio a noviembre de 2026'
                          if is_gantt else 'Árbol de causas y efectos de la situación problemática')
                picture=p.add_run().add_picture(str(FIG_DIR/figure_file),width=width)
                picture._inline.docPr.set('title',alt_text)
                picture._inline.docPr.set('descr',alt_text)
            else:
                for line in content.splitlines():
                    if not line.strip():
                        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); continue
                    leading=len(line)-len(line.lstrip(' '))
                    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.left_indent=Cm(min(3.0,leading*0.12)); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.0
                    set_font(p.add_run(line.strip()),size=10,bold=(leading==0 and not re.match(r'^\d',line.strip())))
        previous_kind=kind

    # Remove the inherited placeholder first paragraph if it is still empty.
    if not doc.paragraphs[0].text and len(doc.paragraphs)>1:
        p=doc.paragraphs[0]._element; p.getparent().remove(p)
    doc.core_properties.title='Sistema de trazabilidad del material vegetal para reforestación con proyección hacia bonos de carbono: caso R3Foresta'
    doc.core_properties.subject='Sistema de trazabilidad del material vegetal para reforestación con proyección hacia bonos de carbono: caso R3Foresta'
    doc.core_properties.author='Pablo Andres Fernandez Cari'
    doc.core_properties.keywords='trazabilidad, cadena de custodia, material vegetal, reforestación, R3Foresta'
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__=='__main__':
    build()
